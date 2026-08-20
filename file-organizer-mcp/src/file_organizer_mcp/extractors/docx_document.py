from __future__ import annotations
from pathlib import Path
from docx import Document
from .common import clean_text, extraction_failure

class DocxExtractor: 
    name = "python-docx" 
    version = "1" 

    def extract(self, path: Path, max_pages: int):
        # Skip temporary Microsoft Word owner files
        if path.name.startswith("~$"):
            return "", {}, extraction_failure("SKIPPED_TEMP_FILE", "Skipping temporary Word lock file.")

        # DOCX does not store reliable rendered page boundaries; MAX_PAGES is not applied.
        del max_pages 
        try: 
            document = Document(str(path))
            parts: list[str] = [] 
            
            for paragraph in document.paragraphs: 
                if paragraph.text: 
                    parts.append(paragraph.text) 
                    
            for table in document.tables: 
                for row in table.rows: 
                    values = [clean_text(cell.text) for cell in row.cells] 
                    parts.append(" | ".join(value for value in values if value)) 
                    
            for section in document.sections: 
                for container in (section.header, section.footer): 
                    for paragraph in container.paragraphs: 
                        if paragraph.text: 
                            parts.append(paragraph.text) 
                            
            text = clean_text("\n".join(parts)) 
            metadata = {"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)} 
            
            if not text: 
                return "", metadata, extraction_failure("NO_EXTRACTABLE_TEXT", "The DOCX file contains no extractable text.") 
                
            return text, metadata, None 
        except Exception as exc: 
            return "", {}, extraction_failure("DOCX_EXTRACTION_FAILED", str(exc))

